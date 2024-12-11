import streamlit as st
from openai import OpenAI
from docx import Document
from docx.shared import Inches
from io import BytesIO
import requests

# Initialize OpenAI client
client = OpenAI(api_key="sk-IyX2YuXfux0xM6UWeOEBhSfdIi-NJz8qadowc8APfGT3BlbkFJqRf90bk4B4WCG0Ql_rIesIQC4OS1y4uqKajygNcv4A")

# CSS for background styling
st.markdown(
    """
    <style>
    .main {
        background-color: #FFA500;
    }
    </style>
    """,
    unsafe_allow_html=True
)

# Function to generate a recipe
def generate_recipe(ingredients, dish_number):
    prompt = (
        f"Create a detailed step-by-step recipe #{dish_number} using these ingredients: {ingredients}. "
        f"Include cooking instructions, total time (preparation + cooking), and tips for better results. "
        f"Also suggest a list of additional ingredients that complement the dish."
    )
    try:
        response = client.chat.completions.create(
            model="gpt-4-turbo",
            messages=[
                {"role": "system", "content": "You are a professional chef."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1000,
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Error generating recipe: {e}"

# Function to generate a recipe image
def generate_recipe_image(dish_name):
    try:
        response = client.images.generate(
            prompt=f"A visually stunning dish called '{dish_name}', served on an elegant plate. Gourmet-style presentation.",
            size="1024x1024"
        )
        return response.data[0].url
    except Exception:
        return "https://via.placeholder.com/1024?text=Image+Not+Available"

# Function to validate ingredients
def validate_ingredients(ingredients):
    required_count = 3  # Minimum number of ingredients
    if len(ingredients.split(',')) < required_count:
        return False, ["Tomato", "Garlic", "Onion"]  # Example additional ingredients
    return True, []

# Function to save recipes to a .docx file
def save_recipes_to_docx(recipes):
    doc = Document()
    doc.add_heading("Generated Recipes", level=1)
    for recipe in recipes:
        doc.add_heading(recipe["title"], level=2)
        doc.add_paragraph(recipe["content"])
        if recipe["image_url"] != "https://via.placeholder.com/1024?text=Image+Not+Available":
            try:
                response = requests.get(recipe["image_url"])
                if response.status_code == 200:
                    doc.add_picture(BytesIO(response.content), width=Inches(4))
            except Exception as e:
                doc.add_paragraph(f"Error adding image: {e}")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit app interface
st.title("AI Recipe Generator 🍴")
ingredients = st.text_input("Enter ingredients (comma-separated):")
num_recipes = st.number_input("How many recipes to generate?", min_value=1, step=1)

if st.button("Generate Recipes"):
    if ingredients:
        is_valid, suggested = validate_ingredients(ingredients)
        if not is_valid:
            st.error(f"Not enough ingredients! Try adding: {', '.join(suggested)}")
        else:
            recipes = []
            for i in range(num_recipes):
                st.info(f"Generating Recipe {i + 1}...")
                recipe_content = generate_recipe(ingredients, i + 1)
                dish_name = recipe_content.split("\n")[0]  # Extract dish name
                image_url = generate_recipe_image(dish_name)
                recipes.append({"title": f"Recipe {i + 1}", "content": recipe_content, "image_url": image_url})
            
            # Display recipes in dynamic columns
            cols = st.columns(num_recipes)
            for i, col in enumerate(cols):
                with col:
                    st.subheader(recipes[i]["title"])
                    st.write(recipes[i]["content"])
                    if recipes[i]["image_url"] != "https://via.placeholder.com/1024?text=Image+Not+Available":
                        st.image(recipes[i]["image_url"], use_column_width=True)

            # Provide a download button for all recipes
            docx_file = save_recipes_to_docx(recipes)
            st.download_button(
                label="Download Recipes",
                data=docx_file,
                file_name="recipes.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.error("Please enter some ingredients!")
